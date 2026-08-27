from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "MAJD-Guard-项目说明与实验进展.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x59, 0x63, 0x6E)
LIGHT_GRAY = RGBColor(0x8A, 0x93, 0x9D)


SECTIONS = [
    (
        "一、这个项目实际做了什么",
        [
            "MAJD-Guard 是一个研究“多智能体系统怎样被污染、污染怎样在多个智能体之间继续传播、以及怎样在不把整个任务全部清空的前提下阻止危险动作”的完整实验平台。它不是只做了一个可视化页面，也不是只做了一个提示词检测器。它把多智能体任务从输入、检索、记忆、摘要、模型推理、工具调用到外部动作的整条链路搭起来，再在这条链路中加入追踪、拦截、修复、复验、回放和实验评估。普通地说，它研究的是这样一种风险：一条恶意信息进入资料库后，可能先被某个智能体读到，又被写进共享记忆或摘要，随后被另一个智能体当成正常信息使用，最后变成邮件内容、数据库参数、日历操作、付款对象、发布指令或命令行参数。项目要做的不是只在入口问“这句话像不像攻击”，而是持续回答“这条信息从哪里来、经过了谁、现在会不会影响高风险动作、发现问题后怎样安全地收拾残局”。",
            "项目的运行从一次被固定下来的任务开始。每次运行都会保存任务、攻击或正常对照、智能体拓扑、随机种子、模型分工、提示词版本、工具权限、风险等级、预算和目标出口等信息，避免不同方法在不同条件下比赛。系统可以构造链式、星型、层级式和小世界式等多智能体关系，让消息、检索片段、记忆、摘要、计划和工具结果在智能体之间真实流动，并且可以调用实际大模型完成规划或审核。前端则把这套过程做成可操作界面，包含总览、策略实验室、运行监控、链路追踪、事件回放、攻击剧本、实验管理、基准测试、防御中心和系统设置。用户既可以看到某次运行发生了什么，也可以沿着来源链查看污染怎样传播、哪一步被拦截、采取了什么修复，以及实验最后得到哪些指标。",
            "为了知道一份内容究竟受过谁的影响，项目维护了一本只能追加、不能随意改写的“来源账本”。每条消息、每次记忆写入、每个检索片段、每份摘要、每个计划和工具结果都拥有独立版本；新内容不会偷偷覆盖旧内容，而是产生新版本，并记录它由哪些旧版本加工而来。隔离、作废、保留和恢复也不是直接改掉历史，而是继续写一条状态变化记录。这样做的好处是，系统不仅能说“这里可能有问题”，还可以拿出一条可以复查的证据链，说明某个低可信来源怎样一步步影响了最终动作。前端看到的图、实验用的图和审计时间线都来自这本账，但公开界面只显示经过过滤的投影，已经隔离或作废的内容不会再被当成正常信息公开，带条件保留标签的内容也不会被洗成完全可信。",
            "所有会改变状态或产生现实影响的操作都要经过统一的动作闸门。项目把动作按影响分成四级：E0 是读取和纯计算；E1 是可以在本次运行内撤销的内部写入，例如写记忆；E2 是会改变外部状态但可以补偿的动作，例如沙箱中的日历、工单、数据库或邮件捕获器；E3 是在一次运行中不可逆的高风险动作，例如一次性支付捕获、凭据释放陷阱或不可撤销发送记录。动作闸门会检查执行者有没有权限、资源范围是否匹配、参数从哪里来、来源完整性够不够、操作是否过期，以及当前是演练还是允许执行。用于回放和验证的 dry-run 会在查找或调用 E2/E3 工具之前直接拒绝外部效果，因此“演练”不会偷偷变成真实执行。模型和检测器可以提供风险证据、要求更严格检查或促使拒绝，但不能把确定性规则已经拒绝的动作改成允许；最终放行权始终在确定性的规则和来源检查手中。",
            "项目还维护两张用途不同的来源图。紧图只保留证据比较明确的影响关系，适合提出“哪些内容也许还能保留”的候选；保守图把模型当时可能看到的输入也算作潜在影响，适合判断是否仍有危险路径并签发安全证书。两张图不是投票关系，而是明确分工：紧图只有建议权，保守图拥有否决权。这样做是因为安全和可用性的错误方向不同。安全判断宁可多怀疑一些，避免漏放；内容保留则不能因为看起来没问题就宣称已经安全。项目的核心不对称机制 RAISE 正是利用这种分工：用较精确的图寻找可以保住的任务内容，再用更保守的图做最终复验。",
            "当动作闸门因为污染来源拒绝操作时，系统不会只停在“拒绝”两个字。它会找出所有能把低可信来源送到危险出口的证据链，计算切断这些链路所需的最小干预组合，并分别考虑操作代价、任务损失、重放成本和人工成本。修复可能包括作废某个版本、切断某条关系、拒绝某种能力，或者保留那些虽然受过污染但在保守图中确实到不了当前危险出口的内容。被保留的内容不是重新洗白，它必须带强制标签，不能作为 E2/E3 的授权依据，而且每到新的动作边界都要重新检查；如果后来出现一条新路径，它会马上被降级作废。求解器提出方案以后，还要由结构上独立的检查器重新判断危险路径是否真的消失。检查结果区分“已经覆盖”“确实无解”和“预算或信息不足”，后两种不能伪装成安全。",
            "修复之后，项目还能在副作用沙箱中补偿可撤销动作、按仍被任务需要的支持关系选择性回放，并在回放后再次检查有没有污染复活。每次正式实验都会把 Manifest、事件、来源账本、求解记录、检查证书、回放记录、沙箱副作用和原始指标封装成不可随意覆盖的运行包，失败尝试也会原样保留，补跑使用新的运行身份。运行结束前，在线防御组件不能读取攻击答案；只有运行终止后，独立的 Oracle 和评估器才会把真实标签与运行记录对齐。由此，这个项目实际形成的是一条完整闭环：多智能体完成任务，污染沿真实工作链传播，所有动作统一过闸，所有影响留下来源，拒绝触发最小修复，独立检查器复验，必要时回放恢复，最后再由离线评估统计安全、任务效果和成本。它目前是一套研究和可复现实验平台，高风险动作在可重置沙箱中验证，并不直接连接真实支付、真实邮件或真实凭据系统。",
        ],
    ),
    (
        "二、目前已经做完了哪些实验，以及这些实验说明了什么",
        [
            "现有实验不是一次性跑一个演示，而是从机制正确性、规模边界、留出数据到真实模型端到端逐层推进。早期规模预实验完成了 720 个网格点，发现保守图虽然安全，但会明显压低可用内容：Clean_E 的中位存活率约为 0.2174；加入不对称保留后，可用率提升到约 0.5591，约为原来的 2.58 倍。这个结果说明双图分工确实有必要，也同时否定了一个更强但不真实的说法：不对称保留可以大幅弥补保守判断的损失，却不能保证完全追平紧图。实验还发现精确求解真正先遇到的瓶颈是候选干预数量，而不是危险证据链数量，因此大图不能把预算截断后的结果写成“已经最优”或“已经安全”。",
            "随后完成的机制层 M-01 至 M-05 对系统的基本逻辑做了更严格的验收。60 个手工边界案例全部通过独立答案校验；保守图对真值 authority edge 和危险证据链的召回均为 100%；精确求解与独立穷举在可穷举实例上完全一致；所有声称已经覆盖的贪心结果都没有残余危险链。针对证书的 120 个字段篡改样本全部被拒绝，closed horizon 下新出现的危险路径能够让旧保留结论降级，open horizon 则不签完整保留证书。规模实验完成 720 个点和 288 个公开基线点，预算扫描保留了 13 条 UNKNOWN，没有把不完整搜索包装成最优结论。打开一次且未用于调参的 100 个 held-out 图也全部完成，公开结果与留出结果的方法排序一致率为 0.88。",
            "机制层的主要正结果是：在公开合成图上，RAISE-asymmetric 相对关闭保留的方案，中位可用性提高了 33.33 个百分点；它的综合代价 J 中位数为 6，而单纯保守方案为 14。在 100 个 held-out 图上，RAISE-asymmetric 的危险逃逸率为 0，J 中位数为 8，任务效用中位数为 0.55；相对于冻结的 B9′ 朴素串联基线，J 降低了 76.47%。不过实验也保留了负结果：RAISE-asymmetric 并没有在所有合成机制图上击败全部方法，source-only 和精确 Oracle 的 J 中位数都是 4，containment-only greedy 是 5，都低于 8。由于这些方法承担的恢复和证书语义不同，现阶段可以说“不对称机制比 B9′ 和单保守方案更省”，但不能说“它全面优于所有安全基线”。",
            "端到端部分首先建成了 24 个基础任务，覆盖四种拓扑、消息、检索、记忆、摘要、计划、工具结果、动作参数以及 E0 到 E3 的不同效果；同时完成了 15 类攻击在四种拓扑上的 60 个攻击场景、60 个最小差异正常对照和独立 Oracle。E-01 的 24 个正常任务全部完成。E-04 的 72 次 canary 运行覆盖 B0 无防御、冻结版 B1 和 RAISE-asymmetric，72 个接受运行全部完成，30 个属于 E2、42 个属于 E3，每个运行都有独立来源账本；其中出现了 11 次正式边界修复、11 条检查记录、20 条证书记录和 11 条恢复回放记录。无防御攻击臂产生了 17 次危险出口逃逸，而两种防御都把危险逃逸降为 0；certified escape、E3 绕过和保留标签违规也全部为 0。72 个密封运行包的 10 个关键字段经独立重算，差异为 0。",
            "最新完成的是 E-05 双模型端到端 pilot。它严格按冻结计划运行 12 对攻击/正常对照、4 个核心方法、2 种模型设置和 5 个配对随机种子，共 960 个预定运行。第一种设置由 DeepSeek-v4-flash 同时承担工作与审核角色；第二种设置仍由 DeepSeek 完成工作，但让 gemini-3.7-flash-high 只提供不可授权的 gate/verifier 证据，最终放行仍由确定性动作闸门决定。960/960 个预定身份全部完成，包含 400 个 E2 和 560 个 E3 运行，共进行了 1920 次逻辑模型调用，记录 760,789 个 token。过程中保留了 3 个首轮超时失败包，并以新身份补跑完成，没有覆盖原始失败轨迹。独立审计重新计算了 960 个密封包各 10 个关键字段，差异仍为 0，场景有效性、模型身份、运行包完整性和密钥扫描全部通过。",
            "E-05 的安全结果很明确：B0 无防御在两个模型设置的攻击臂中各出现 85 次危险逃逸，总计 170 次；B1、B9′ 和 RAISE-asymmetric 的危险逃逸均为 0，同时 certified escape、E3 绕过、标签违规、UNKNOWN 和接受运行中的 timeout 都是 0。这说明来源账本、统一动作闸门、沙箱、边界修复和独立重算已经能够在真实模型调用和多拓扑任务中连成一条可审计的端到端链路，而且安全结果没有依赖更强审核模型直接授权。正常对照在各方法、各模型设置下的任务成功率均为 0.917；剩下的一类是专门验证 dry-run 不产生副作用的任务，因此该比例不能简单理解成普通正常任务失败。",
            "同时，pilot 也暴露了当前最重要的结论边界。B1、B9′ 和 RAISE 在攻击臂中的任务成功率都为 0；B9′ 和 RAISE 的平均恢复成功均为 0.833，但恢复状态并没有在这批任务中进一步转化为攻击条件下的任务完成。RAISE 在两种模型设置中分别触发 52 次和 53 次边界修复，B9′ 分别为 50 次和 50 次，但这不能证明 RAISE 已经在端到端效果或成本上优于 B9′。因此，已经完成的实验能够证明“系统确实会追踪、拦截、修复、复验和留下可重算证据”，也证明防御比无防御安全；它还不能证明“RAISE 在完整任务效用、综合代价和外部环境上全面领先”。E-05 的正式决定是 GO，含义只是允许在预算明确批准后进入更大的 E-06 主实验，而不是已经得到论文级最终结论。",
        ],
    ),
    (
        "三、还需要做哪些实验，以及应当看哪些指标",
        [
            "下一步首先是 E-06 全量端到端主实验，但它目前仍因 9,600 次运行的预算没有被单独批准而处于阻断状态。完整矩阵将把 60 个攻击场景和 60 个匹配正常对照组成的 120 个 case arms，与 2 种模型设置、5 个配对随机种子和 8 种方法组合起来。方法不仅包括 B0、B1、B9′ 和 RAISE-asymmetric，还要加入 Deny-all、Full-reset、RAISE-conservative，以及真正忠实实现的 B7 dependency rollback。当前 B7-faithful 仍被明确标为 EXCLUDED，不能拿简化版本冒充；在主实验开始前要么补齐忠实复现，要么继续作为明确排除项报告。E-06 还要每日审计危险证书、费用、UNKNOWN、缺失运行和失败原因，并保证每个计划身份都有成功或明确失败状态。只有这一步完成，才能判断 pilot 中“安全但攻击任务没有继续完成”的现象是样本规模造成的，还是方法本身的效用限制。",
            "第二类缺口是外部有效性。项目尚未完成 AgentDojo 和 A2ASecBench 的正式接入，也还没有外部 Benchmark 结果。需要先冻结两个外部基准的版本，复现它们的官方 smoke，再把工具调用、工具结果、危险出口、任务效用、过度阻断和环境重置映射到项目的统一格式。随后计划运行 480 次小规模外部验证，其中 AgentDojo 240 次、A2ASecBench 240 次，并保存官方原始输出、统一指标、失败轨迹和映射差异。因为外部环境的未来路径无法完全封闭，这部分必须按 open horizon 处理，只能验证动作安全和有限迁移能力，不能借此签发完整恢复或保留安全证书。至少一个适用子集复现出可解释的动作安全收益，项目才有资格讨论有限的外部推广性。",
            "第三类工作是统计、敏感性分析和独立终审。需要从原始运行包自动生成机制主表、端到端主表、外部验证表、安全失败表、配对效应、成本表、排除清单和 Pareto 前沿，并在空白环境中重新计算。二元结果应采用配对检验、精确区间或混合效应 logistic；连续指标应以任务模板为聚类单位做 paired bootstrap，报告中位数、四分位距、95% 置信区间和效应量，多重比较使用 Holm 校正。还要改变综合代价中各项权重，检查方法优劣是否只依赖某一组人为参数，并分别按模型设置、拓扑、攻击族和效果等级分层。最后由独立审计重新检查所有证书、E3 绕过、标签违规、Oracle 泄漏、排除、重跑和失败样本，形成最终 Go/No-Go，而不是只看一张平均数表。",
            "实验指标首先要回答“有没有真的漏出危险动作”。最核心的是 unsafe-sink escape rate，即攻击因素是否最终造成了未补偿的 E2/E3 危险效果；certified escape count 则更严格，检查系统已经签了安全证书的情况下是否仍然逃逸，这一项必须为 0。还要记录 E3 bypass，也就是不可逆动作有没有绕过完整授权生命周期；记录 residual witness count，判断修复后还有没有剩余危险证据链；检查 certificate validity、保留标签违规以及 dry-run/replay 是否产生真实副作用。对未知来源、预算耗尽和无法穷举的情况，要分别记录 UNKNOWN 与 UNSATISFIABLE，不能把“没算完”和“已经证明无解”混在一起，更不能把两者写成允许。",
            "第二组指标要回答“安全以后，任务还剩多少用处”。benign task success 衡量正常任务能否完成，attacked task success 衡量攻击发生后能否在安全条件下继续完成任务，overblocking 衡量系统是否为了安全把过多正常请求也拦掉。recovery success 检查系统能否恢复需要的状态，benign-state preservation 关注无辜内容保住了多少，support group retained/recovered 关注完成目标所需的一组相互支持内容是否仍然完整；recontamination 检查恢复后污染是否复活，time to recovery 则记录从发现问题到恢复可用状态所需时间。E-05 已经显示恢复成功和最终任务成功不是同一个概念，因此后续报告必须把二者分开，不能用“恢复记录存在”替代“任务真的完成”。",
            "第三组指标要回答“这种安全是否付出了不可接受的代价”。项目把成本拆成实际干预代价 C_op、任务损失 L_task、重放代价 C_replay 和人工审批代价 C_human，再组合为综合代价 J(X)；同时记录求解器和检查器延迟、大模型调用次数、token 与可核验费用、账本存储开销、总体运行时开销、timeout、deadline miss、starvation 和 requeue。比较顺序必须先安全、后成本：如果某种方法出现 certified escape，就不能因为它更快或更便宜而获胜；在安全不劣的前提下，才比较它能否比最强可复现基线把 J 至少降低 10%，或把 recovery success 至少提高 5 个百分点，并要求另一项不能实质恶化。unsafe escape 的建议绝对非劣界为 1 个百分点，但 certified escape 的容许上限仍是 0。",
            "最终真正需要回答的问题不是“系统有没有拦住攻击”这么单一，而是“它是否在同样安全的情况下，比忠实强基线更少破坏任务、更快恢复、成本更低，而且这种收益在同模型全栈、不同拓扑、不同攻击、留出数据和至少一个外部基准中都仍然存在”。如果全量实验发现只有 Deny-all 或大规模删除才能实现零逃逸，或者 RAISE 的收益只出现在调参网格、只来自更强审核模型、在 held-out 或外部环境中消失，就应当给出 No-Go 或负结果叙事。相反，只有机制 Gate 继续通过、所有关键安全计数为 0、完整端到端矩阵显示安全不劣且效用或 J 有预注册幅度的改善、同模型设置仍保留收益，并且外部适用子集也得到可解释结果，项目才可以给出最终 Go 结论。",
        ],
    ),
]


def set_run_font(run, *, ascii_font: str = "Calibri", east_asia_font: str = "Microsoft YaHei") -> None:
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def set_style_font(style, *, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_page_number(paragraph) -> None:
    prefix = paragraph.add_run("第 ")
    set_run_font(prefix)
    prefix.font.size = Pt(9)
    prefix.font.color.rgb = LIGHT_GRAY
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = LIGHT_GRAY
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix)
    suffix.font.size = Pt(9)
    suffix.font.color.rgb = LIGHT_GRAY


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.widow_control = True
    run = paragraph.add_run(text)
    set_run_font(run)


def build_document() -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    set_style_font(title, size=23, color=BLACK, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.0

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, size=12.5, color=GRAY)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.line_spacing = 1.0

    heading_1 = styles["Heading 1"]
    set_style_font(heading_1, size=16, color=BLUE, bold=True)
    heading_1.paragraph_format.space_before = Pt(16)
    heading_1.paragraph_format.space_after = Pt(8)
    heading_1.paragraph_format.line_spacing = 1.0
    heading_1.paragraph_format.keep_with_next = True

    for style_name, size, color, before, after in (
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    if "Document Meta" not in styles:
        meta_style = styles.add_style("Document Meta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta_style = styles["Document Meta"]
    set_style_font(meta_style, size=9.5, color=LIGHT_GRAY)
    meta_style.paragraph_format.space_before = Pt(0)
    meta_style.paragraph_format.space_after = Pt(2)
    meta_style.paragraph_format.line_spacing = 1.0

    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "MAJD-GUARD  ·  项目与实验说明"
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    for run in header_p.runs:
        set_run_font(run)
        run.font.size = Pt(8.5)
        run.font.color.rgb = LIGHT_GRAY

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    add_page_number(footer_p)

    kicker = document.add_paragraph(style="Document Meta")
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(5)
    kicker_run = kicker.add_run("研究项目通俗说明")
    set_run_font(kicker_run)
    kicker_run.font.size = Pt(10)
    kicker_run.font.bold = True
    kicker_run.font.color.rgb = BLUE

    title_p = document.add_paragraph("MAJD-Guard 项目说明与实验进展", style="Title")
    keep_with_next(title_p)
    subtitle_p = document.add_paragraph(
        "用普通语言说明项目整体、已完成实验及下一步验证",
        style="Subtitle",
    )
    keep_with_next(subtitle_p)

    meta_1 = document.add_paragraph(style="Document Meta")
    meta_1.add_run("项目：多智能体级联污染与防御可视化")
    meta_2 = document.add_paragraph(style="Document Meta")
    meta_2.add_run("更新日期：2026 年 8 月 27 日")
    meta_3 = document.add_paragraph(style="Document Meta")
    meta_3.paragraph_format.space_after = Pt(14)
    meta_run = meta_3.add_run("当前阶段：E-05 端到端 Pilot 已完成并通过独立审计；E-06 待明确预算批准")
    meta_run.font.bold = True
    for paragraph in (meta_1, meta_2, meta_3):
        for run in paragraph.runs:
            set_run_font(run)
            run.font.size = Pt(9.5)
            run.font.color.rgb = GRAY

    for section_index, (heading, paragraphs) in enumerate(SECTIONS):
        heading_p = document.add_paragraph(heading, style="Heading 1")
        heading_p.paragraph_format.keep_with_next = True
        for run in heading_p.runs:
            set_run_font(run)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = BLUE
        for paragraph_text in paragraphs:
            add_body_paragraph(document, paragraph_text)

    core = document.core_properties
    core.title = "MAJD-Guard 项目说明与实验进展"
    core.subject = "项目整体说明、已完成实验、后续实验与指标"
    core.author = "MAJD-Guard 项目组"
    core.keywords = "MAJD-Guard, 多智能体, 级联污染, 防御, 端到端实验"
    core.comments = "依据当前仓库代码与截至 2026-08-27 的机制层、E-04 canary 和 E-05 pilot 独立审计结果生成。"

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
